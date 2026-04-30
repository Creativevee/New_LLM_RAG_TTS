# rag.py - Fully Optimized Version
import os
import re
import time
from datetime import datetime, timezone
from pathlib import Path
from huggingface_hub import InferenceClient
from pypdf import PdfReader
import chromadb
from chromadb.utils import embedding_functions
import edge_tts
import tempfile
import json
import csv

# --- CONFIG ---
from dotenv import load_dotenv
load_dotenv()
HF_TOKEN = os.environ.get("HF_TOKEN", "")
MODEL_ID = "Qwen/Qwen2.5-7B-Instruct"
VOICE = "en-US-GuyNeural"
ALLOWED_EXTENSIONS = {'.pdf', '.txt', '.docx', '.csv', '.json', '.md'}

# Initialize
client = InferenceClient(token=HF_TOKEN)
chroma_client = chromadb.Client()

#keeping memory of chats
EFE = embedding_functions.DefaultEmbeddingFunction()
collection = chroma_client.get_or_create_collection(
    name="knowledge_base",
    embedding_function=EFE
)

#creating another collection for storing chat memory
chat_collection = chroma_client.get_or_create_collection(
    name="chat_memory",
    embedding_function=EFE
)

#storing chats
def store_chat_message(session_id: str, role: str, text: str):
    # Take out the symbol because it was giving encoding errors on Windows terminals.
    print(f"Storing chat message: session_id={session_id}, role={role}, text={repr(text[:50])}...")
    chat_collection.add(
        ids=[f"{session_id}_{int(time.time() * 1e6)}"],
        documents=[text],
        metadatas=[{"session_id": session_id, "role": role}]
    )

def get_recent_chat_history(session_id: str, limit=10):
    results = chat_collection.get(where={"session_id": session_id})
    items = [
        {"role": m["role"], "text": d, "_id": id_}
        for d, m, id_ in zip(
            results["documents"], results["metadatas"], results["ids"]
        )
        if m.get("type") != "thread_meta"
    ]
    items.sort(key=lambda x: x["_id"])
    return [{"role": i["role"], "text": i["text"]} for i in items[-limit:]]

#creating threads:
def create_thread(session_id: str, title: str = "Untitled"):
    #ensuring thread exists, storing thread metadata
    first_message = "System: Thread created." + f" Title: {title}"
    chat_collection.add(
        ids=[f"{session_id}_0"],
        documents=[first_message],
        metadatas=[{
            "session_id": session_id,
            "role": "system",
            "type": "thread_meta",
            "title": title,
            "created_at": datetime.now(timezone.utc).isoformat()
        }]
    )

def get_thread_list():
    #fetching all distinct session_ids(threads)
    results = chat_collection.get(
        where={"type": "thread_meta"}
    )
    threads = []
    for m in results["metadatas"]:
        threads.append({
            "session_id": m["session_id"],
            "title": m.get("title", "Untitled"),
            "created_at": m["created_at"],
        })
    return threads


def reset_thread_messages(session_id: str) -> int:
    """Delete user/assistant messages for a thread, keep thread metadata."""
    results = chat_collection.get(where={"session_id": session_id})
    ids = results.get("ids", []) or []
    metadatas = results.get("metadatas", []) or []

    ids_to_delete = []
    for msg_id, metadata in zip(ids, metadatas):
        metadata = metadata or {}
        if metadata.get("type") != "thread_meta":
            ids_to_delete.append(msg_id)

    if ids_to_delete:
        chat_collection.delete(ids=ids_to_delete)

    return len(ids_to_delete)

# --- TEXT CLEANING ---
def clean_text(text):
    if not text:
        return ""
    text = re.sub(r'\r\n', '\n', text)       # normalise Windows line endings
    text = re.sub(r'\r', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)   # collapse 3+ blank lines to 2
    text = re.sub(r'[^\S\n]+', ' ', text)    # collapse horizontal whitespace only, keep \n
    return text.strip()

# --- DOCUMENT LOADERS ---
def load_pdf(path):
    reader = PdfReader(path)
    text = ""
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text += f"\n\n[PAGE {i+1}]\n{page_text}"
    return text, {"source": path.name, "pages": len(reader.pages), "type": ".pdf"}

def load_docx(path):
    from docx import Document
    doc = Document(path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return text, {"source": path.name, "pages": len(doc.paragraphs), "type": ".docx"}

def load_txt(path):
    encodings = ['utf-8', 'latin-1', 'cp1252']
    for enc in encodings:
        try:
            text = path.read_text(encoding=enc)
            return text, {"source": path.name, "pages": 1, "type": ".txt"}
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Could not decode {path.name}")

def load_csv(path):
    with open(path, 'r', encoding='utf-8') as f:
        reader = csv.reader(f)
        rows = list(reader)
        text = "\n".join([", ".join(row) for row in rows])
    return text, {"source": path.name, "pages": len(rows), "type": ".csv"}

def load_json(path):
    with open(path, 'r', encoding='utf-8') as f:
        data = json.load(f)
        text = json.dumps(data, indent=2)
    return text, {"source": path.name, "pages": 1, "type": ".json"}

def load_document(file_path):
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File {file_path} not found")

    loaders = {
        '.pdf': load_pdf,
        '.txt': load_txt,
        '.md': load_txt,
        '.docx': load_docx,
        '.csv': load_csv,
        '.json': load_json
    }

    suffix = path.suffix.lower()
    if suffix not in loaders:
        raise ValueError(f"Unsupported file type: {suffix}")

    text, metadata = loaders[suffix](path)
    text = clean_text(text)
    metadata["char_count"] = len(text)
    return text, metadata

def validate_file(filename: str) -> bool:
    ext = Path(filename).suffix.lower()
    return ext in ALLOWED_EXTENSIONS

# --- SMART CHUNKING ---
def split_text(text, chunk_size=1000, chunk_overlap=150):
    # Break into paragraphs; sub-split long ones into sentences
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
    units = []
    for para in paragraphs:
        if len(para) <= chunk_size:
            units.append(para)
        else:
            sentences = re.split(r'(?<=[.!?])\s+', para)
            units.extend(s.strip() for s in sentences if s.strip())

    chunks = []
    current = ""
    for unit in units:
        if not current:
            current = unit
        elif len(current) + len(unit) + 2 <= chunk_size:
            current += "\n\n" + unit
        else:
            chunks.append(current)
            # carry the tail of the previous chunk into the next for overlap
            overlap = current[-chunk_overlap:].lstrip()
            space = overlap.find(' ')
            if space > 0:
                overlap = overlap[space + 1:]
            current = (overlap + "\n\n" + unit) if overlap else unit

    if current:
        chunks.append(current)
    return chunks

# --- INDEXING ---
def index_document(file_content: bytes, filename: str):
    start_time = time.time()

    with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp:
        tmp.write(file_content)
        tmp_path = tmp.name

    try:
        text, metadata = load_document(tmp_path)
        metadata["source"] = filename
        chunks = split_text(text)

        try:
            existing = collection.get(where={"source": filename})
            if existing["ids"]:
                collection.delete(ids=existing["ids"])
        except:
            pass

        safe_name = re.sub(r'[^\w.-]', '_', filename)
        batch_size = 50
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i+batch_size]
            collection.add(
                ids=[f"{safe_name}_{i+j}" for j in range(len(batch))],
                documents=batch,
                metadatas=[metadata for _ in batch]
            )

        elapsed = time.time() - start_time
        return len(chunks), elapsed, metadata
    finally:
        os.unlink(tmp_path)

# --- SEARCH ---
def get_all_context():
    """Return every chunk in the knowledge base. Used for 'list all' type questions
    where similarity search would miss entries that score low but are still relevant."""
    results = collection.get()
    if not results["ids"]:
        return "", []
    documents = results.get("documents") or []
    metadatas = results.get("metadatas") or []
    if not documents:
        return "", []
    context = "\n\n".join(d for d in documents if d)
    seen = {}
    for meta in metadatas:
        if not meta:
            continue
        src = meta.get("source", "uploaded_document")
        if src not in seen:
            seen[src] = meta.get("type", "unknown")
    sources = [{"source": s, "type": t, "chunk_index": 1, "preview": ""}
               for s, t in seen.items()]
    return context, sources

def search_context(query, n_results=10):
    count = collection.count()
    if count == 0:
        return "", []
    results = collection.query(query_texts=[query], n_results=min(n_results, count))
    documents = results.get("documents", [[]])[0]
    metadatas = results.get("metadatas", [[]])[0]

    if not documents:
        return "", []

    context = "\n\n".join(documents)
    sources = []
    for idx, doc in enumerate(documents):
        metadata = metadatas[idx] if idx < len(metadatas) and metadatas[idx] else {}
        source_name = metadata.get("source", "uploaded_document")
        source_type = metadata.get("type", "unknown")
        preview = doc[:220].strip()
        if len(doc) > 220:
            preview += "..."

        sources.append(
            {
                "source": source_name,
                "type": source_type,
                "chunk_index": idx + 1,
                "preview": preview,
            }
        )

    return context, sources

# --- LLM ---
def get_answer(question, context, history=None):
    #Combining conversation history with context if available
    context_parts = []
    if history and history.strip():
        context_parts.append(f"CHAT HISTORY:\n{history}")
    if context and context.strip() and context != "No relevant document context found":
        context_parts.append(f"DOCUMENT CONTEXT:\n{context}")
    if context_parts:
        context_block = "\n\n" + "\n\n".join(context_parts)
        system_prompt = (
            "You are a friendly and knowledgeable AI assistant — think of how ChatGPT answers questions: "
            "naturally, conversationally, and helpfully. Use the document context as your primary source of facts, "
            "but express the answer in your own words. Do not copy sentences verbatim from the document.\n\n"
            "Important rules:\n"
            "- When a question asks you to list people, roles, or items, list every single one separately — "
            "never merge two people into one bullet point.\n"
            "- Write in a natural, friendly tone. Synthesise the information rather than quoting it.\n"
            "- If the context covers the answer fully, use it. If the question goes beyond the document, "
            "blend your own knowledge in naturally without saying 'the document says'.\n"
            "- Be complete. For list questions, do not stop early or skip entries."
        )
        user_content = f"{context_block}\n\nQUESTION: {question}\n\nANSWER:"
    else:
        system_prompt = (
            "You are a friendly and knowledgeable AI assistant. Answer conversationally and helpfully, "
            "like ChatGPT would. Be clear, natural, and complete."
        )
        user_content = question
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_content}
    ]
    try:
        response = client.chat.completions.create(
            model=MODEL_ID,
            messages=messages,
            max_tokens=700,
            temperature=0.4
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error: {e}"

# --- TTS ---
async def text_to_speech(text):
    audio_data = b""
    communicate = edge_tts.Communicate(text, VOICE)
    async for chunk in communicate.stream():
        if chunk["type"] == "audio":
            audio_data += chunk["data"]
    return audio_data

# --- ADMIN: KB MANAGEMENT ---
def get_kb_documents():
    results = collection.get()
    if not results["ids"]:
        return []
    sources = {}
    for meta in (results.get("metadatas") or []):
        if not meta:
            continue
        src = meta.get("source", "unknown")
        if src not in sources:
            sources[src] = {
                "source": src,
                "type": meta.get("type", "unknown"),
                "chunks": 0,
            }
        sources[src]["chunks"] += 1
    return list(sources.values())

def remove_document_from_kb(source_name: str) -> int:
    results = collection.get(where={"source": source_name})
    if results["ids"]:
        collection.delete(ids=results["ids"])
    return len(results["ids"])



